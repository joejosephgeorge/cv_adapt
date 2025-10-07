"""
Utility functions for CV Adaptor
Document parsing, web scraping, and formatting utilities
"""
from typing import Optional
import os


def extract_text_from_file(file_path: str, filename: str) -> Optional[str]:
    """
    Extract text from various document formats
    
    Args:
        file_path: Path to the file
        filename: Original filename
    
    Returns:
        Extracted text or None if failed
    """
    file_ext = filename.split('.')[-1].lower()
    
    try:
        if file_ext == 'pdf':
            return _extract_from_pdf(file_path)
        elif file_ext == 'docx':
            return _extract_from_docx(file_path)
        elif file_ext == 'txt':
            return _extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    except Exception as e:
        print(f"Error extracting text: {str(e)}")
        return None


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF"""
    try:
        import pypdf
        
        text = []
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text())
        
        return "\n".join(text)
    except ImportError:
        # Fallback to PyPDF2
        try:
            import PyPDF2
            
            text = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            
            return "\n".join(text)
        except:
            raise ImportError("Neither pypdf nor PyPDF2 is installed. Install with: pip install pypdf")


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return "\n".join(text)
    except ImportError:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")


def _extract_from_txt(file_path: str) -> str:
    """Extract text from TXT"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    raise ValueError("Unable to decode text file with supported encodings")


def scrape_job_description(url: str) -> Optional[str]:
    """
    Scrape job description from URL
    
    Args:
        url: Job posting URL
    
    Returns:
        Extracted job description text or None
    """
    try:
        import requests
        from bs4 import BeautifulSoup
        
        # Set headers to mimic browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "header", "footer", "nav"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    except ImportError:
        raise ImportError("Required packages not installed. Install with: pip install requests beautifulsoup4")
    except Exception as e:
        print(f"Error scraping URL: {str(e)}")
        return None
